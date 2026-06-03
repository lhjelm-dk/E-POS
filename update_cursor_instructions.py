#!/usr/bin/env python3
"""
Append the latest changes addendum to GeoRisk_CursorInstructions.docx.
Requires: pip install python-docx
Usage: python update_cursor_instructions.py
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

DOCX_PATH = Path(r"c:\Users\lhjel\Downloads\GeoRisk_CursorInstructions.docx")
OUTPUT_PATH = Path(r"c:\Users\lhjel\Downloads\GeoRisk_CursorInstructions_updated.docx")

ADDENDUM = """
GeoRisk Cursor Instructions — Update Addendum (March 2025)

Summary of Implemented Changes

1. Shared Risk Element Model
All three methods (Classic POS, ESL, Bayesian) now use the same risk elements: Play level (4 pillars) + Conditional level (sub-elements). Sub-elements combined with min (weakest link).

2. Play Chance and Conditional Prospect Probability
All methods display: Play Chance, Conditional Prospect Chance, Total Pg/POS/P(Discovery) = Play × Conditional.

3. Bayesian Network — Full Risk Element Model
Bayesian tab now has Play Chance, Conditional sub-elements, and blue overview table with Play + Conditional layout.

4. Blue Overview Tables
All overview tables use consistent blue styling (#eff6ff, #3b82f6).

5. Sub-Element Combination — Min
Clear captions: "min = weakest link. Sub-elements are not independent."

6. Combination Hierarchy — Unified Structure
All three hierarchy charts show: Result ← Play × Conditional ← 4 pillars each. Method-specific combination rules in footer.

7. Method Differentiation
Only combination differs: Classic POS (product + min), ESL (product + ALL/ANY/IPT), Bayesian (product via BN + min).

Key files: app.py, methods/classic_pos.py, methods/bayesian.py, components/overview_table.py, components/hierarchy_chart.py
"""


def main():
    if not DOCX_PATH.exists():
        print(f"Source file not found: {DOCX_PATH}")
        return
    doc = Document(str(DOCX_PATH))
    # Add new page and addendum
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GeoRisk Cursor Instructions — Update Addendum")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("March 2025")
    doc.add_paragraph()
    for line in ADDENDUM.strip().split("\n\n"):
        if line.startswith("Summary") or line.startswith("Key files"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line and line[0].isdigit() and "." in line[:3]:
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(str(OUTPUT_PATH))
    print(f"Updated document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
